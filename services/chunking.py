"""
Извлечение текста из распространённых форматов и разбиение на чанки для RAG.

Опциональные зависимости (ставьте по мере нужды):
    pip install pypdf python-docx openpyxl

Текстовые файлы читаются как UTF-8 с заменой некорректных байтов (errors=\"replace\").
"""

from __future__ import annotations

import csv
import json
import re
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

# --- форматы без внешних зависимостей ---
_TEXT_EXTENSIONS = frozenset({
    ".txt",
    ".md",
    ".rst",
    ".log",
    ".csv",
    ".tsv",
    ".json",
    ".jsonl",
    ".xml",
    ".yaml",
    ".yml",
    ".toml",
    ".ini",
    ".cfg",
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".java",
    ".go",
    ".rs",
    ".c",
    ".h",
    ".cpp",
    ".hpp",
    ".cs",
    ".rb",
    ".php",
    ".swift",
    ".kt",
    ".sql",
    ".sh",
    ".bash",
    ".zsh",
    ".env",
})
_HTML_EXTENSIONS = frozenset({".html", ".htm", ".xhtml"})


class _HTMLTextExtractor(HTMLParser):
    """Сбор видимого текста из HTML без BeautifulSoup."""

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._chunks: list[str] = []

    def handle_data(self, data: str) -> None:
        if data.strip():
            self._chunks.append(data)

    def get_text(self) -> str:
        raw = "".join(self._chunks)
        raw = re.sub(r"[ \t]+\n", "\n", raw)
        return re.sub(r"\n{3,}", "\n\n", raw)


def _normalize_ws(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _read_plain(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _read_csv_like(path: Path, delimiter: str | None = None) -> str:
    lines: list[str] = []
    with path.open(newline="", encoding="utf-8", errors="replace") as f:
        sample = f.read(4096)
        f.seek(0)
        if delimiter is None:
            try:
                dialect = csv.Sniffer().sniff(sample, delimiters=",;\t")
                delimiter = dialect.delimiter
            except csv.Error:
                delimiter = ","
        reader = csv.reader(f, delimiter=delimiter)
        for row in reader:
            lines.append("\t".join(cell.strip() for cell in row))
    return "\n".join(lines)


def _flatten_json(obj: Any, prefix: str = "") -> list[str]:
    out: list[str] = []
    if isinstance(obj, dict):
        for k, v in obj.items():
            key = f"{prefix}.{k}" if prefix else str(k)
            out.extend(_flatten_json(v, key))
    elif isinstance(obj, list):
        for i, v in enumerate(obj):
            out.extend(_flatten_json(v, f"{prefix}[{i}]"))
    else:
        out.append(f"{prefix}: {obj}" if prefix else str(obj))
    return out


def _read_json(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="replace")
    data = json.loads(raw)
    if isinstance(data, list) and data and isinstance(data[0], dict):
        # JSON Lines часто хранят как один массив или построчно — здесь обычный JSON
        pass
    lines = _flatten_json(data)
    return "\n".join(lines)


def _read_jsonl(path: Path) -> str:
    lines: list[str] = []
    with path.open(encoding="utf-8", errors="replace") as f:
        for line_no, line in enumerate(f, start=1):
            line = line.strip()
            if not line:
                continue
            try:
                obj = json.loads(line)
            except json.JSONDecodeError:
                lines.append(line)
                continue
            flat = _flatten_json(obj, prefix=f"row{line_no}")
            lines.extend(flat)
    return "\n".join(lines)


def _read_html(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="replace")
    parser = _HTMLTextExtractor()
    parser.feed(raw)
    return _normalize_ws(parser.get_text())


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ImportError(
            'Для PDF установите: pip install pypdf'
        ) from e
    reader = PdfReader(str(path))
    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        if t.strip():
            parts.append(t)
    return _normalize_ws("\n\n".join(parts))


def _read_docx(path: Path) -> str:
    try:
        import docx  # type: ignore
    except ImportError as e:
        raise ImportError(
            'Для DOCX установите: pip install python-docx'
        ) from e
    document = docx.Document(str(path))
    blocks: list[str] = []
    for p in document.paragraphs:
        if p.text.strip():
            blocks.append(p.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append("\t".join(cells))
    return _normalize_ws("\n\n".join(blocks))


def _read_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as e:
        raise ImportError(
            'Для XLSX установите: pip install openpyxl'
        ) from e
    wb = load_workbook(str(path), read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in wb.worksheets:
        lines.append(f"## {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(x.strip() for x in cells):
                lines.append("\t".join(cells))
    wb.close()
    return _normalize_ws("\n".join(lines))


def extract_text(path: str | Path) -> tuple[str, dict[str, Any]]:
    """
    Читает файл и возвращает (текст, базовые метаданные).

    Метаданные: source_path, suffix, extractor.
    """
    p = Path(path).expanduser().resolve()
    if not p.is_file():
        raise FileNotFoundError(p)
    suf = p.suffix.lower()
    meta: dict[str, Any] = {"source_path": str(p), "suffix": suf, "extractor": ""}

    if suf == ".csv" or suf == ".tsv":
        meta["extractor"] = "csv"
        text = _read_csv_like(p, delimiter="\t" if suf == ".tsv" else None)
    elif suf == ".jsonl" or suf == ".ndjson":
        meta["extractor"] = "jsonl"
        text = _read_jsonl(p)
    elif suf == ".json":
        meta["extractor"] = "json"
        text = _read_json(p)
    elif suf in _HTML_EXTENSIONS:
        meta["extractor"] = "html"
        text = _read_html(p)
    elif suf == ".pdf":
        meta["extractor"] = "pdf"
        text = _read_pdf(p)
    elif suf == ".docx":
        meta["extractor"] = "docx"
        text = _read_docx(p)
    elif suf in {".xlsx", ".xlsm"}:
        meta["extractor"] = "xlsx"
        text = _read_xlsx(p)
    elif suf in _TEXT_EXTENSIONS or not suf:
        meta["extractor"] = "plain"
        text = _read_plain(p)
    else:
        meta["extractor"] = "plain_fallback"
        text = _read_plain(p)

    return _normalize_ws(text), meta


@dataclass
class TextChunk:
    """Один чанк для последующего эмбеддинга."""

    text: str
    metadata: dict[str, Any] = field(default_factory=dict)


def _split_oversized_piece(text: str, max_chars: int, overlap: int) -> list[str]:
    """Жёсткое окно по символам для очень длинных абзацев."""
    if max_chars <= 0:
        raise ValueError("max_chars должен быть > 0")
    if overlap < 0 or overlap >= max_chars:
        raise ValueError("overlap должен быть >= 0 и < max_chars")
    step = max_chars - overlap
    out: list[str] = []
    start = 0
    n = len(text)
    while start < n:
        out.append(text[start : start + max_chars])
        start += step
    return out


def chunk_text(
    text: str,
    max_chars: int = 1200,
    overlap: int = 150,
    *,
    respect_paragraphs: bool = True,
) -> list[str]:
    """
    Разбивает текст на чанки.

    - Сначала режем по абзацам (двойной перевод строки), пока помещается в max_chars.
    - Если абзац длиннее max_chars — режем скользящим окном с перекрытием overlap.
    """
    text = _normalize_ws(text)
    if not text:
        return []

    if not respect_paragraphs:
        return _split_oversized_piece(text, max_chars, overlap)

    paragraphs = [x.strip() for x in re.split(r"\n\s*\n+", text) if x.strip()]
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0
    sep_len = 2

    def flush() -> None:
        nonlocal current, current_len
        if current:
            chunks.append("\n\n".join(current))
            current = []
            current_len = 0

    for para in paragraphs:
        plen = len(para)
        if plen > max_chars:
            flush()
            chunks.extend(_split_oversized_piece(para, max_chars, overlap))
            continue
        add = plen + (sep_len if current else 0)
        if current_len + add <= max_chars:
            current.append(para)
            current_len += add
        else:
            flush()
            current.append(para)
            current_len = plen

    flush()
    return chunks


def chunk_file(
    path: str | Path,
    max_chars: int = 1200,
    overlap: int = 150,
    *,
    respect_paragraphs: bool = True,
    extra_metadata: dict[str, Any] | None = None,
) -> list[TextChunk]:
    """
    Извлекает текст из файла и возвращает список TextChunk.

    В metadata каждого чанка: source_path, suffix, extractor, chunk_index, chunk_count,
    плюс необязательный extra_metadata (одинаковый merge в каждый чанк).
    """
    text, base_meta = extract_text(path)
    if not text:
        return []

    strings = chunk_text(
        text,
        max_chars=max_chars,
        overlap=overlap,
        respect_paragraphs=respect_paragraphs,
    )
    total = len(strings)
    extra = extra_metadata or {}
    out: list[TextChunk] = []
    for i, s in enumerate(strings):
        meta = {
            **base_meta,
            **extra,
            "chunk_index": i,
            "chunk_count": total,
        }
        out.append(TextChunk(text=s, metadata=meta))
    return out


def chunk_paths(
    paths: list[str | Path],
    max_chars: int = 1200,
    overlap: int = 150,
    *,
    respect_paragraphs: bool = True,
) -> list[TextChunk]:
    """Несколько файлов подряд; глобальные индексы по всему списку."""
    all_chunks: list[TextChunk] = []
    for path in paths:
        all_chunks.extend(
            chunk_file(
                path,
                max_chars=max_chars,
                overlap=overlap,
                respect_paragraphs=respect_paragraphs,
            )
        )
    total = len(all_chunks)
    for i, c in enumerate(all_chunks):
        c.metadata["global_chunk_index"] = i
        c.metadata["global_chunk_count"] = total
    return all_chunks
